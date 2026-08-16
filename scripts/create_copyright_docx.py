# -*- coding: utf-8 -*-
"""Create software copyright registration document for Pioneer AI Service Framework V2.0.0"""

import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_name="\u5b8b\u4f53", font_size=12, bold=False):
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_heading_styled(doc, text, level=1, font_size=15):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "\u9ed1\u4f53", font_size, bold=True)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body_text(doc, text, indent_first=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "\u5b8b\u4f53", 12)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_bullet_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "\u5b8b\u4f53", 12)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    return p


def set_cell_text(cell, text, font_name="\u5b8b\u4f53", font_size=12, bold=False, align_center=False):
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, font_name, font_size, bold)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_column_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)

    # ========== COVER PAGE ==========
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u8ba1\u7b97\u673a\u8f6f\u4ef6\u8457\u4f5c\u6743\u767b\u8bb0\u7533\u8bf7")
    set_run_font(run, "\u9ed1\u4f53", 22, bold=True)

    doc.add_paragraph()

    cover_info = [
        ("\u8f6f\u4ef6\u540d\u79f0\uff1a\u5148\u950b\u4eba\u5de5\u667a\u80fd\u670d\u52a1\u6846\u67b6\u8f6f\u4ef6", 16),
        ("\uff08Pioneer AI Service Framework\uff09", 14),
        ("\u7248\u672c\u53f7\uff1aV2.0.0", 14),
    ]
    for text, fs in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "\u5b8b\u4f53", fs)

    for _ in range(4):
        doc.add_paragraph()

    items = [
        ("\u8457\u4f5c\u6743\u4eba", "\u79e6\u6653\u671b"),
        ("\u5f00\u53d1\u5b8c\u6210\u65e5\u671f", "2026\u5e7405\u6708"),
        ("\u9996\u6b21\u53d1\u8868\u65e5", "\u672a\u53d1\u8868"),
        ("\u8f6f\u4ef6\u5206\u7c7b", "\u5e94\u7528\u8f6f\u4ef6"),
        ("\u7b80\u79f0", "\u5148\u950bAI\u670d\u52a1\u6846\u67b6"),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}\uff1a{value}")
        set_run_font(run, "\u5b8b\u4f53", 14)

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading_styled(doc, "\u76ee  \u5f55", level=0, font_size=18)
    toc = [
        ("\u4e00\u3001\u8f6f\u4ef6\u6982\u51b5", 14, True),
        ("    1.1 \u8f6f\u4ef6\u57fa\u672c\u4fe1\u606f", 12, False),
        ("    1.2 \u8f6f\u4ef6\u529f\u80fd\u6982\u8ff0", 12, False),
        ("    1.3 \u6280\u672f\u67b6\u6784", 12, False),
        ("\u4e8c\u3001\u8f6f\u4ef6\u6280\u672f\u7279\u70b9", 14, True),
        ("    2.1 \u6838\u5fc3\u529f\u80fd\u6a21\u5757", 12, False),
        ("    2.2 \u6280\u672f\u5b9e\u73b0\u8def\u5f84", 12, False),
        ("    2.3 V2.0.0 \u65b0\u589e\u529f\u80fd", 12, False),
        ("\u4e09\u3001\u7f16\u7a0b\u8bed\u8a00\u4e0e\u8fd0\u884c\u73af\u5883", 14, True),
        ("    3.1 \u7f16\u7a0b\u8bed\u8a00", 12, False),
        ("    3.2 \u8fd0\u884c\u73af\u5883", 12, False),
        ("    3.3 \u6570\u636e\u5e93", 12, False),
        ("\u56db\u3001\u8f6f\u4ef6\u6587\u6863\u6e05\u5355", 14, True),
    ]
    for text, fs, bold in toc:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, "\u5b8b\u4f53", fs, bold)
        p.paragraph_format.line_spacing = Pt(26)

    doc.add_page_break()

    # ========== SECTION 1 ==========
    add_heading_styled(doc, "\u4e00\u3001\u8f6f\u4ef6\u6982\u51b5", level=1, font_size=16)

    add_heading_styled(doc, "1.1 \u8f6f\u4ef6\u57fa\u672c\u4fe1\u606f", level=2, font_size=14)
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    info_data = [
        ("\u8f6f\u4ef6\u540d\u79f0", "\u5148\u950b\u4eba\u5de5\u667a\u80fd\u670d\u52a1\u6846\u67b6\u8f6f\u4ef6\uff08Pioneer AI Service Framework\uff09"),
        ("\u7248\u672c\u53f7", "V2.0.0"),
        ("\u8457\u4f5c\u6743\u4eba", "\u79e6\u6653\u671b"),
        ("\u5f00\u53d1\u5b8c\u6210\u65e5\u671f", "2026\u5e7405\u6708"),
        ("\u9996\u6b21\u53d1\u8868\u65e5", "\u672a\u53d1\u8868"),
        ("\u8f6f\u4ef6\u5206\u7c7b", "\u5e94\u7528\u8f6f\u4ef6"),
        ("\u7b80\u79f0", "\u5148\u950bAI\u670d\u52a1\u6846\u67b6"),
    ]
    for i, (k, v) in enumerate(info_data):
        set_cell_text(table.rows[i].cells[0], k, "\u5b8b\u4f53", 12, bold=True, align_center=True)
        set_cell_text(table.rows[i].cells[1], v, "\u5b8b\u4f53", 12)
        set_column_width(table.rows[i].cells[0], 4000)
        set_column_width(table.rows[i].cells[1], 11000)

    doc.add_paragraph()

    # 1.2 Overview
    add_heading_styled(doc, "1.2 \u8f6f\u4ef6\u529f\u80fd\u6982\u8ff0", level=2, font_size=14)
    add_body_text(doc, "\u5148\u950b\u4eba\u5de5\u667a\u80fd\u670d\u52a1\u6846\u67b6\u8f6f\u4ef6\uff08Pioneer AI Service Framework\uff09\u662f\u4e00\u5957\u9762\u5411 Pi Network \u751f\u6001\u7684\u767d\u6807\u5546\u6237\u5e94\u7528\u6a21\u677f\u6846\u67b6\u3002")
    add_body_text(doc, "\u7cfb\u7edf\u91c7\u7528 Monorepo \u67b6\u6784\uff08pnpm + Turborepo\uff09\uff0c\u4ee5 Next.js 14 App Router + TypeScript \u4e3a\u6838\u5fc3\u6280\u672f\u6808\uff0c\u5185\u7f6e\u4f01\u4e1a\u7ea7 AI \u670d\u52a1\u8def\u7531\u5f15\u64ce\uff0c\u652f\u6301 OpenAI / Anthropic / Ollama \u591a\u6a21\u578b\u52a8\u6001\u5207\u6362\u4e0e\u81ea\u52a8\u5bb9\u9519\u3002")
    add_body_text(doc, "\u7cfb\u7edf\u8bbe\u8ba1\u7406\u5ff5\u4e3a\u201c80% \u901a\u7528\u5e95\u5ea7 + 20% \u884c\u4e1a\u914d\u7f6e\u201d\uff0c\u901a\u8fc7\u884c\u4e1a\u9884\u8bbe\u4e0e\u5546\u6237\u914d\u7f6e\u673a\u5236\uff0c\u53ef\u5728\u6781\u77ed\u5468\u671f\u5185\u4ea4\u4ed8\u5b9a\u5236\u5316\u5546\u6237\u5e94\u7528\uff0c\u8986\u76d6\u7f8e\u5bb9\u7f8e\u7532\u3001\u5065\u8eab\u3001\u57f9\u8bad\u3001\u54a8\u8be2\u7b49\u591a\u79cd\u5782\u76f4\u884c\u4e1a\u3002")

    features = [
        "\u591a AI \u63d0\u4f9b\u5546\u667a\u80fd\u8def\u7531\uff1a\u57fa\u4e8e Strategy + Factory \u8bbe\u8ba1\u6a21\u5f0f\uff0c\u652f\u6301 OpenAI\u3001Anthropic (Claude)\u3001Ollama \u4e09\u5927 AI \u63d0\u4f9b\u5546",
        "Pi Network \u539f\u751f\u652f\u4ed8\uff1a\u5b8c\u6574\u96c6\u6210 Pi U2A\uff08User-to-App\uff09\u652f\u4ed8\u6d41\u7a0b",
        "\u591a\u79df\u6237\u67b6\u6784\uff1aV2.0.0 \u65b0\u589e\u5546\u6237\u7ea7\u6570\u636e\u786c\u9694\u79bb",
        "License \u6388\u6743\u9a8c\u8bc1\uff1a\u57fa\u4e8e HMAC-SHA256 \u7684\u79bb\u7ebf\u6388\u6743\u9a8c\u8bc1",
        "\u7528\u91cf\u7edf\u8ba1\u4e0e\u914d\u989d\u7ba1\u7406\uff1a\u5fae\u79d2\u7ea7 API \u8c03\u7528\u4e0e Token \u6d88\u8017\u8ffd\u8e2a",
        "\u5b89\u5168\u8ba4\u8bc1\u4f53\u7cfb\uff1aPi SDK \u8ba4\u8bc1 + HMAC \u7b7e\u540d Session + HttpOnly Cookie",
        "\u884c\u4e1a\u914d\u7f6e\u5305\uff1a\u901a\u8fc7\u7ed3\u6784\u5316\u914d\u7f6e\u9a71\u52a8 UI \u6e32\u67d3\u3001\u529f\u80fd\u6a21\u5757\u5f00\u5173\u4e0e\u4e1a\u52a1\u6d41\u7a0b\u5b9a\u5236",
    ]
    for f in features:
        add_bullet_text(doc, f"\u2022  {f}")

    # 1.3 Architecture
    add_heading_styled(doc, "1.3 \u6280\u672f\u67b6\u6784", level=2, font_size=14)
    add_body_text(doc, "\u672c\u7cfb\u7edf\u91c7\u7528\u524d\u540e\u7aef\u540c\u6784\u7684\u5168\u6808\u67b6\u6784\uff0c\u7531 Next.js 14 App Router \u7edf\u4e00\u7ba1\u7406\u524d\u7aef\u6e32\u67d3\u4e0e\u670d\u52a1\u7aef\u903b\u8f91\u3002\u4ee3\u7801\u4ed3\u5e93\u91c7\u7528 Monorepo \u7ed3\u6784\uff0c\u901a\u8fc7 pnpm Workspace \u4e0e Turborepo \u5b9e\u73b0\u9ad8\u6548\u7684\u5305\u7ba1\u7406\u4e0e\u6784\u5efa\u7f13\u5b58\u3002")

    table2 = doc.add_table(rows=9, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Table Grid"

    arch_headers = ["\u5c42\u7ea7", "\u6280\u672f\u9009\u578b", "\u8bf4\u660e"]
    arch_data = [
        ("\u524d\u7aef\u6846\u67b6", "Next.js 14 (App Router)", "React Server Components + \u6d41\u5f0f\u6e32\u67d3"),
        ("\u7f16\u7a0b\u8bed\u8a00", "TypeScript 5.4+", "\u4e25\u683c\u6a21\u5f0f\uff0c\u5168\u9879\u76ee\u7c7b\u578b\u5b89\u5168"),
        ("UI \u6837\u5f0f", "Tailwind CSS 3", "\u539f\u5b50\u5316 CSS\uff0c\u6309\u9700\u7f16\u8bd1"),
        ("\u540e\u7aef", "Next.js API Routes", "\u540c\u6784\u670d\u52a1\u7aef\uff0c\u65e0\u9700\u72ec\u7acb\u540e\u7aef"),
        ("\u6570\u636e\u5e93", "PostgreSQL 15+", "\u5173\u7cfb\u578b\u6570\u636e\u5e93\uff0cPrisma ORM \u6620\u5c04"),
        ("ORM", "Prisma 5", "\u7c7b\u578b\u5b89\u5168\u7684\u6570\u636e\u5e93\u8bbf\u95ee\u5c42"),
        ("\u5305\u7ba1\u7406", "pnpm 8 (Monorepo)", "\u9ad8\u6548\u4f9d\u8d56\u89e3\u6790\u4e0e\u78c1\u76d8\u5229\u7528"),
        ("AI \u5f15\u64ce", "\u591a\u63d0\u4f9b\u5546\u8def\u7531\u5668", "OpenAI / Anthropic / Ollama \u667a\u80fd\u5207\u6362"),
    ]
    for j, h in enumerate(arch_headers):
        set_cell_text(table2.rows[0].cells[j], h, "\u9ed1\u4f53", 11, bold=True, align_center=True)
    for i, (a, b, c) in enumerate(arch_data):
        vals = [a, b, c]
        for j, val in enumerate(vals):
            set_cell_text(table2.rows[i+1].cells[j], val, "\u5b8b\u4f53", 11, align_center=(j == 0))
    for row in table2.rows:
        for j, cell in enumerate(row.cells):
            set_column_width(cell, [3000, 4000, 8000][j])

    doc.add_page_break()

    # ========== SECTION 2 ==========
    add_heading_styled(doc, "\u4e8c\u3001\u8f6f\u4ef6\u6280\u672f\u7279\u70b9", level=1, font_size=16)

    add_heading_styled(doc, "2.1 \u6838\u5fc3\u529f\u80fd\u6a21\u5757", level=2, font_size=14)
    modules = [
        ("AI \u667a\u80fd\u8def\u7531\u6a21\u5757",
         "\u652f\u6301 OpenAI GPT-4o-mini\u3001Anthropic Claude\u3001Ollama \u591a\u6a21\u578b\u63a5\u5165\uff0c\u57fa\u4e8e Strategy + Factory \u6a21\u5f0f\u5b9e\u73b0\u52a8\u6001\u5207\u6362\u4e0e\u81ea\u52a8\u5bb9\u9519\u964d\u7ea7\u3002\u5f53\u4e3b\u63d0\u4f9b\u5546\u5931\u8d25\u65f6\uff0c\u7cfb\u7edf\u6309\u914d\u7f6e\u987a\u5e8f\u81ea\u52a8\u5c1d\u8bd5\u5907\u7528\u63d0\u4f9b\u5546\u3002"),
        ("Pi Network \u652f\u4ed8\u96c6\u6210",
         "\u5b8c\u6574\u5b9e\u73b0 Pi U2A \u652f\u4ed8\u6d41\u7a0b\uff0c\u5305\u62ec\u521b\u5efa\u652f\u4ed8\u3001\u5ba1\u6838\u3001\u94fe\u4e0a\u786e\u8ba4\u3001\u5b8c\u6210\u7684\u5b8c\u6574\u751f\u547d\u5468\u671f\u7ba1\u7406\u3002\u6240\u6709\u652f\u4ed8\u64cd\u4f5c\u5747\u5728\u670d\u52a1\u7aef\u6267\u884c\u3002"),
        ("\u5546\u6237\u914d\u7f6e\u5316\u5f15\u64ce",
         "\u901a\u8fc7\u7ed3\u6784\u5316\u914d\u7f6e\u6587\u4ef6\u9a71\u52a8 UI \u6e32\u67d3\u3001\u529f\u80fd\u6a21\u5757\u5f00\u5173\u4e0e\u4e1a\u52a1\u6d41\u7a0b\u5b9a\u5236\u3002\u652f\u6301\u7f8e\u5bb9\u7f8e\u7532\u3001\u5065\u8eab\u3001\u57f9\u8bad\u3001\u54a8\u8be2\u7b49\u591a\u79cd\u884c\u4e1a\u9884\u8bbe\u3002"),
        ("\u5b89\u5168\u8ba4\u8bc1\u4f53\u7cfb",
         "\u91c7\u7528\u4e09\u5c42\u8ba4\u8bc1\u673a\u5236\uff1aPi SDK \u7aef\u8ba4\u8bc1 \u2192 Pi Platform API \u4e8c\u6b21\u9a8c\u8bc1 \u2192 HMAC \u7b7e\u540d Session\u3002Session \u5b58\u50a8\u5728 HttpOnly + Secure Cookie \u4e2d\uff0c\u6709\u6548\u671f 7 \u5929\u3002"),
        ("\u591a\u79df\u6237\u67b6\u6784 (V2.0.0)",
         "\u5546\u6237\u7ea7\u6570\u636e\u786c\u9694\u79bb\uff0c\u57fa\u4e8e AsyncLocalStorage \u6ce8\u5165\u79df\u6237\u4e0a\u4e0b\u6587\uff0cPrisma \u4e2d\u95f4\u4ef6\u81ea\u52a8\u6ce8\u5165 merchantId \u8fc7\u6ee4\u6761\u4ef6\u3002"),
        ("License \u6388\u6743\u8ba4\u8bc1 (V2.0.0)",
         "\u91c7\u7528 HMAC-SHA256 \u7684\u79bb\u7ebf\u6388\u6743\u9a8c\u8bc1\u673a\u5236\uff0c\u652f\u6301 Starter / Professional / Enterprise \u4e09\u7ea7\u5957\u9910\u63a7\u5236\u3002\u5f00\u53d1\u73af\u5883\u4e0b\u81ea\u52a8\u9881\u53d1\u4f01\u4e1a\u7ea7\u8bb8\u53ef\u3002"),
    ]
    for title, desc in modules:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}\uff1a")
        set_run_font(run, "\u9ed1\u4f53", 12, bold=True)
        run2 = p.add_run(desc)
        set_run_font(run2, "\u5b8b\u4f53", 12)
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0.74)

    # 2.2 Technical paths
    add_heading_styled(doc, "2.2 \u6280\u672f\u5b9e\u73b0\u8def\u5f84", level=2, font_size=14)
    add_body_text(doc, "\u672c\u7cfb\u7edf\u57fa\u4e8e TypeScript \u5168\u6808\u5f00\u53d1\uff0c\u91c7\u7528\u4ee5\u4e0b\u5173\u952e\u6280\u672f\u8def\u5f84\uff1a")
    paths = [
        "Monorepo \u5de5\u7a0b\u7ba1\u7406\uff1apnpm Workspace + Turborepo\uff0c\u786e\u4fdd\u591a\u5305\u534f\u8c03\u5f00\u53d1\u4e0e\u9ad8\u6548\u6784\u5efa",
        "React Server Components\uff1aNext.js 14 App Router \u539f\u751f\u652f\u6301\uff0c\u670d\u52a1\u7aef\u6e32\u67d3\u51cf\u5c11\u5ba2\u6237\u7aef JS \u4f53\u79ef",
        "Prisma ORM\uff1a\u7c7b\u578b\u5b89\u5168\u7684\u6570\u636e\u8bbf\u95ee\u5c42\uff0c\u81ea\u52a8\u751f\u6210 TypeScript \u7c7b\u578b\u5b9a\u4e49",
        "Pi SDK\uff1a\u5c01\u88c5 Pi Network \u8ba4\u8bc1\u4e0e\u652f\u4ed8\u63a5\u53e3\uff0c\u63d0\u4f9b\u7edf\u4e00\u8c03\u7528\u62bd\u8c61",
        "AI \u591a\u63d0\u4f9b\u5546\u8def\u7531\uff1aStrategy \u6a21\u5f0f\u5b9e\u73b0\u53ef\u63d2\u62d4 AI \u63d0\u4f9b\u5546\u652f\u6301",
        "\u6d41\u5f0f\u54cd\u5e94\uff1aServer-Sent Events (SSE) \u5b9e\u73b0 AI \u54cd\u5e94\u7684\u5b9e\u65f6\u6d41\u5f0f\u8f93\u51fa",
    ]
    for path in paths:
        add_bullet_text(doc, f"\u2022  {path}")

    # 2.3 V2.0.0 new features
    add_heading_styled(doc, "2.3 V2.0.0 \u65b0\u589e\u529f\u80fd", level=2, font_size=14)
    add_body_text(doc, "\u76f8\u6bd4 V1.0.0 \u7248\u672c\uff0cV2.0.0 \u65b0\u589e\u4ee5\u4e0b\u6838\u5fc3\u529f\u80fd\uff1a")
    new_features = [
        "\u591a\u79df\u6237\u67b6\u6784\uff1a\u5b9e\u73b0\u5546\u6237\u7ea7\u6570\u636e\u786c\u9694\u79bb\uff0c\u652f\u6301\u5355\u4e00\u6846\u67b6\u5b9e\u4f8b\u670d\u52a1\u6d77\u91cf\u5546\u6237",
        "License \u6388\u6743\u9a8c\u8bc1\u7cfb\u7edf\uff1aHMAC-SHA256 \u79bb\u7ebf\u6388\u6743\uff0c\u652f\u6301\u4e09\u7ea7\u5546\u4e1a\u6388\u6743\u63a7\u5236",
        "\u7528\u91cf\u7edf\u8ba1\u4e0e\u914d\u989d\u7ba1\u7406\uff1a\u5fae\u79d2\u7ea7 API \u8c03\u7528\u4e0e Token \u6d88\u8017\u8ffd\u8e2a\uff0c\u6708\u5ea6\u8ba2\u9605\u5236\u8ba1\u8d39\u6a21\u578b",
        "\u589e\u5f3a\u578b AI \u8def\u7531\uff1a\u65b0\u589e Anthropic Claude \u63d0\u4f9b\u5546\u652f\u6301\uff0c\u5b8c\u5584\u81ea\u52a8\u5bb9\u9519\u964d\u7ea7\u673a\u5236",
        "\u884c\u4e1a\u914d\u7f6e\u5305\u5347\u7ea7\uff1a\u6269\u5c55\u66f4\u591a\u5782\u76f4\u884c\u4e1a\u9884\u8bbe\u4e0e\u914d\u7f6e\u6a21\u677f",
    ]
    for nf in new_features:
        add_bullet_text(doc, f"\u2022  {nf}")

    doc.add_page_break()

    # ========== SECTION 3 ==========
    add_heading_styled(doc, "\u4e09\u3001\u7f16\u7a0b\u8bed\u8a00\u4e0e\u8fd0\u884c\u73af\u5883", level=1, font_size=16)

    add_heading_styled(doc, "3.1 \u7f16\u7a0b\u8bed\u8a00", level=2, font_size=14)
    languages = [
        "TypeScript 5.4+\uff08\u4e25\u683c\u6a21\u5f0f\uff0c\u5168\u9879\u76ee\u7c7b\u578b\u5b89\u5168\uff09\u2014 \u6838\u5fc3\u5f00\u53d1\u8bed\u8a00",
        "JavaScript (ES2022+) \u2014 \u8f85\u52a9\u811a\u672c\u4e0e\u914d\u7f6e",
        "SQL (PostgreSQL) \u2014 \u6570\u636e\u5e93\u67e5\u8be2\u4e0e\u8fc1\u79fb",
        "HTML5 / CSS3 / Tailwind CSS \u2014 \u524d\u7aef\u8868\u73b0\u5c42",
        "Python 3.10+ \u2014 \u8f85\u52a9\u5de5\u5177\u811a\u672c",
    ]
    for lang in languages:
        add_bullet_text(doc, f"\u2022  {lang}")

    add_heading_styled(doc, "3.2 \u8fd0\u884c\u73af\u5883", level=2, font_size=14)
    envs = [
        "\u8fd0\u884c\u65f6\uff1aNode.js >= 18.0.0",
        "\u64cd\u4f5c\u7cfb\u7edf\uff1aWindows 10+ / macOS 12+ / Ubuntu 20.04+",
        "\u5305\u7ba1\u7406\uff1apnpm >= 8.0.0",
        "AI \u6a21\u578b\u63a5\u53e3\uff1aOpenAI API / Anthropic API / Ollama \u672c\u5730",
        "\u533a\u5757\u94fe\u7f51\u7edc\uff1aPi Network Mainnet",
        "\u6700\u4f4e\u5185\u5b58\uff1a4 GB",
        "\u63a8\u8350\u914d\u7f6e\uff1a8 GB \u5185\u5b58 + SSD",
    ]
    for env in envs:
        add_bullet_text(doc, f"\u2022  {env}")

    add_heading_styled(doc, "3.3 \u6570\u636e\u5e93", level=2, font_size=14)
    add_body_text(doc, "\u7cfb\u7edf\u4f7f\u7528 PostgreSQL 15+ \u4f5c\u4e3a\u4e3b\u8981\u6570\u636e\u5b58\u50a8\uff0c\u901a\u8fc7 Prisma ORM \u8fdb\u884c\u6570\u636e\u5e93\u8bbf\u95ee\u4e0e\u8fc1\u79fb\u7ba1\u7406\u3002")

    db_tables = [
        ("merchants", "\u5546\u6237\u57fa\u672c\u4fe1\u606f"),
        ("customers", "\u987e\u5ba2\u6863\u6848"),
        ("services", "\u670d\u52a1\u76ee\u5f55"),
        ("orders", "\u8ba2\u5355\u8bb0\u5f55"),
        ("payments", "Pi \u652f\u4ed8\u8bb0\u5f55"),
        ("memberships", "\u4f1a\u5458\u65b9\u6848"),
        ("bookings", "\u9884\u7ea6\u8bb0\u5f55"),
        ("api_usage", "API \u7528\u91cf\u7edf\u8ba1\uff08V2.0.0 \u65b0\u589e\uff09"),
        ("licenses", "\u6388\u6743\u8bc1\u4e66\u8bb0\u5f55\uff08V2.0.0 \u65b0\u589e\uff09"),
    ]
    table3 = doc.add_table(rows=len(db_tables) + 1, cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = "Table Grid"
    set_cell_text(table3.rows[0].cells[0], "\u8868\u540d", "\u9ed1\u4f53", 11, bold=True, align_center=True)
    set_cell_text(table3.rows[0].cells[1], "\u804c\u8d23", "\u9ed1\u4f53", 11, bold=True, align_center=True)
    set_column_width(table3.rows[0].cells[0], 5000)
    set_column_width(table3.rows[0].cells[1], 10000)
    for i, (tbl, desc) in enumerate(db_tables):
        set_cell_text(table3.rows[i + 1].cells[0], tbl, "\u5b8b\u4f53", 11, align_center=True)
        set_cell_text(table3.rows[i + 1].cells[1], desc, "\u5b8b\u4f53", 11)
        set_column_width(table3.rows[i + 1].cells[0], 5000)
        set_column_width(table3.rows[i + 1].cells[1], 10000)

    doc.add_page_break()

    # ========== SECTION 4 ==========
    add_heading_styled(doc, "\u56db\u3001\u8f6f\u4ef6\u6587\u6863\u6e05\u5355", level=1, font_size=16)
    add_body_text(doc, "\u672c\u8f6f\u4ef6\u8457\u4f5c\u6743\u767b\u8bb0\u7533\u8bf7\u63d0\u4ea4\u4ee5\u4e0b\u914d\u5957\u6587\u6863\uff1a")

    docs_list = [
        ("01-\u6e90\u4ee3\u7801\u6587\u6863-\u5148\u950bAI\u670d\u52a1\u6846\u67b6V2.0.0.pdf", "\u6e90\u4ee3\u7801\u6587\u6863\uff08\u516060\u9875\uff09", "\u5305\u542b\u6838\u5fc3\u6e90\u4ee3\u7801\u5c55\u793a"),
        ("02-\u7528\u6237\u64cd\u4f5c\u624b\u518c-\u5148\u950bAI\u670d\u52a1\u6846\u67b6V2.0.0.pdf", "\u7528\u6237\u64cd\u4f5c\u624b\u518c", "\u7cfb\u7edf\u529f\u80fd\u4ecb\u7ecd\u4e0e\u64cd\u4f5c\u6307\u5357"),
    ]
    table4 = doc.add_table(rows=len(docs_list) + 1, cols=3)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.style = "Table Grid"
    table_headers = ["\u6587\u4ef6\u540d", "\u5185\u5bb9\u8bf4\u660e", "\u5907\u6ce8"]
    for j, h in enumerate(table_headers):
        set_cell_text(table4.rows[0].cells[j], h, "\u9ed1\u4f53", 11, bold=True, align_center=True)
        set_column_width(table4.rows[0].cells[j], [5000, 5000, 5000][j])
    for i, (fname, desc, note) in enumerate(docs_list):
        vals = [fname, desc, note]
        for j, val in enumerate(vals):
            set_cell_text(table4.rows[i + 1].cells[j], val, "\u5b8b\u4f53", 10)
            set_column_width(table4.rows[i + 1].cells[j], [5000, 5000, 5000][j])

    doc.add_paragraph()
    add_body_text(doc, "\u4ee5\u4e0a\u6587\u6863\u5171\u540c\u6784\u6210\u5148\u950b\u4eba\u5de5\u667a\u80fd\u670d\u52a1\u6846\u67b6\u8f6f\u4ef6 V2.0.0 \u7684\u5b8c\u6574\u8457\u4f5c\u6743\u767b\u8bb0\u7533\u8bf7\u6750\u6599\u3002\u6e90\u4ee3\u7801\u6587\u6863\u63d0\u4f9b\u4e86\u6309\u9875\u5206\u5272\u7684\u6838\u5fc3\u4ee3\u7801\u5c55\u793a\uff0c\u7528\u6237\u64cd\u4f5c\u624b\u518c\u63d0\u4f9b\u4e86\u5b8c\u6574\u7684\u7cfb\u7edf\u529f\u80fd\u8bf4\u660e\u4e0e\u64cd\u4f5c\u6307\u5357\u3002")

    # Save
    output_path = r"D:\PiMerchantFramework\先锋人工智能服务框架软件-软著登记申请-V2.0.0.docx"
    doc.save(output_path)
    print(f"[SUCCESS] \u8f6f\u8457\u767b\u8bb0\u7533\u8bf7\u6587\u6863\u5df2\u751f\u6210: {output_path}")
    print(f"         \u6587\u4ef6\u5927\u5c0f: {os.path.getsize(output_path):,} \u5b57\u8282")


if __name__ == "__main__":
    build_document()
